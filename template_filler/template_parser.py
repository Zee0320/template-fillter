"""
TemplateParser: DOCX 模板解析与填充

解析 Word 文档中的 {{KEYWORD}} 占位符，并在保留格式的情况下替换内容。
"""

import re
from docx import Document
from typing import Dict, List, Tuple


class TemplateParser:
    """DOCX 模板解析器"""
    
    PLACEHOLDER_PATTERN = re.compile(r'\{\{(\w+)\}\}')
    
    def __init__(self, template_path: str):
        """
        初始化模板解析器
        
        Args:
            template_path: Word 模板文件路径
        """
        self.template_path = template_path
        self.document = Document(template_path)
    
    def find_placeholders(self) -> List[str]:
        """
        查找模板中所有的占位符
        
        Returns:
            占位符名称列表（不包含双花括号）
        """
        placeholders = set()
        
        # 遍历所有段落
        for paragraph in self.document.paragraphs:
            text = paragraph.text
            matches = self.PLACEHOLDER_PATTERN.findall(text)
            placeholders.update(matches)
        
        # 遍历所有表格
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        matches = self.PLACEHOLDER_PATTERN.findall(text)
                        placeholders.update(matches)
        
        return list(placeholders)
    
    def fill_placeholders(self, content_map: Dict[str, str]) -> None:
        """
        填充占位符，保留原有格式
        
        Args:
            content_map: 占位符名称 -> 填充内容 的映射
        """
        # 处理段落
        for paragraph in self.document.paragraphs:
            self._replace_in_paragraph(paragraph, content_map)
        
        # 处理表格
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, content_map)
    
    def _replace_in_paragraph(self, paragraph, content_map: Dict[str, str]) -> None:
        """
        在段落中替换占位符，保留格式
        
        策略：在 run 级别进行替换，以保留字体、颜色等格式
        """
        # 获取段落的完整文本
        full_text = paragraph.text
        
        # 检查是否包含占位符
        if not self.PLACEHOLDER_PATTERN.search(full_text):
            return
        
        # 对于每个占位符，查找并替换
        for placeholder, content in content_map.items():
            pattern = f'{{{{{placeholder}}}}}'
            if pattern in full_text:
                # 遍历 runs 查找包含占位符的 run
                self._replace_in_runs(paragraph.runs, pattern, content)
    
    def _replace_in_runs(self, runs, pattern: str, content: str) -> None:
        """
        在 runs 中替换占位符
        
        处理跨 run 的占位符情况（如 "{{" 在一个 run，"NAME}}" 在另一个 run）
        """
        # 简单情况：占位符完全在一个 run 内
        for run in runs:
            if pattern in run.text:
                run.text = run.text.replace(pattern, content)
                return
        
        # 复杂情况：占位符跨多个 runs
        # 合并所有 run 的文本，找到占位符位置，然后重建
        combined_text = ''.join(run.text for run in runs)
        if pattern not in combined_text:
            return
        
        # 计算每个 run 的字符范围
        run_ranges = []
        pos = 0
        for run in runs:
            run_ranges.append((pos, pos + len(run.text), run))
            pos += len(run.text)
        
        # 找到占位符的位置
        start_idx = combined_text.find(pattern)
        end_idx = start_idx + len(pattern)
        
        # 找到涉及的 runs
        affected_runs = []
        for start, end, run in run_ranges:
            if start < end_idx and end > start_idx:
                affected_runs.append((start, end, run))
        
        if not affected_runs:
            return
        
        # 在第一个受影响的 run 中放置替换内容，清空其他受影响的 runs
        first_run = affected_runs[0][2]
        first_start = affected_runs[0][0]
        
        # 计算需要保留的前缀和后缀
        prefix = combined_text[:start_idx]
        suffix = combined_text[end_idx:]
        new_text = prefix + content + suffix
        
        # 重建 runs
        char_pos = 0
        for start, end, run in run_ranges:
            run_len = end - start
            if char_pos + run_len <= len(new_text):
                run.text = new_text[char_pos:char_pos + run_len]
            else:
                run.text = new_text[char_pos:] if char_pos < len(new_text) else ''
            char_pos += run_len
    
    def save(self, output_path: str) -> None:
        """
        保存填充后的文档
        
        Args:
            output_path: 输出文件路径
        """
        self.document.save(output_path)


if __name__ == '__main__':
    # 简单测试
    import sys
    if len(sys.argv) > 1:
        parser = TemplateParser(sys.argv[1])
        placeholders = parser.find_placeholders()
        print(f"Found placeholders: {placeholders}")
