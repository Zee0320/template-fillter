"""
创建测试模板 DOCX 文件

由于无法直接创建 Word 文件，使用 python-docx 生成包含各种格式的测试模板。
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os


def create_test_template():
    """创建包含各种格式的测试模板"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('', 0)
    title_run = title.add_run('项目报告：{{TITLE}}')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空行
    doc.add_paragraph()
    
    # 摘要部分
    summary_heading = doc.add_heading('摘要', level=1)
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('{{SUMMARY}}')
    summary_run.font.size = Pt(12)
    summary_run.font.italic = True
    
    # 空行
    doc.add_paragraph()
    
    # 关键词部分
    keywords_para = doc.add_paragraph()
    keywords_label = keywords_para.add_run('关键词：')
    keywords_label.font.bold = True
    keywords_label.font.size = Pt(11)
    keywords_content = keywords_para.add_run('{{KEYWORDS}}')
    keywords_content.font.size = Pt(11)
    keywords_content.font.color.rgb = RGBColor(0, 102, 153)
    
    # 空行
    doc.add_paragraph()
    
    # 项目意义部分
    significance_heading = doc.add_heading('项目意义', level=1)
    significance_para = doc.add_paragraph()
    significance_run = significance_para.add_run('{{SIGNIFICANCE}}')
    significance_run.font.size = Pt(12)
    
    # 创建表格测试
    doc.add_heading('项目概览', level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # 表格内容
    cells = table.rows[0].cells
    cells[0].text = '项目标题'
    cells[1].text = '{{TITLE}}'
    
    cells = table.rows[1].cells
    cells[0].text = '核心价值'
    cells[1].text = '{{SIGNIFICANCE}}'
    
    cells = table.rows[2].cells
    cells[0].text = '关键词'
    cells[1].text = '{{KEYWORDS}}'
    
    # 保存
    output_path = os.path.join(os.path.dirname(__file__), 'test_template.docx')
    doc.save(output_path)
    print(f"测试模板已创建: {output_path}")
    return output_path


if __name__ == '__main__':
    create_test_template()
